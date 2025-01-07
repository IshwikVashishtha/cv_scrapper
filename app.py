import pdfplumber , traceback , os , re , spacy , nltk 
import pandas as pd
from docx import Document
from nameparser import HumanName

# Download required NLTK data
# nltk.download('averaged_perceptron_tagger_eng')
# nltk.download('maxent_ne_chunker_tab')
# nltk.download('words')
# nltk.download('punkt')

# Load spaCy model
nlp = spacy.load('en_core_web_sm')

def extract_text_from_pdf(pdf_path):
    print(f"Extracting text from PDF: {pdf_path}")
    text_by_page = []
    try:
        with pdfplumber.open(pdf_path) as pdf:
            print(f"PDF opened successfully, pages: {len(pdf.pages)}")
            for page_num, page in enumerate(pdf.pages, 1):
                text = page.extract_text()
                print(f"Page {page_num} extracted, length: {len(text) if text else 0}")
                text_by_page.append(text or '')
    except Exception as e:
        print(f"Error in PDF extraction: {str(e)}")
        print(traceback.format_exc())
    return text_by_page

def extract_text_from_doc(doc_path):
    print(f"Extracting text from DOC: {doc_path}")
    text_by_page = []
    current_page = []
    
    try:
        doc = Document(doc_path)
        print(f"Document opened successfully, paragraphs: {len(doc.paragraphs)}")
        
        for para_num, paragraph in enumerate(doc.paragraphs, 1):
            if paragraph.text.strip():
                current_page.append(paragraph.text)
                print(f"Added paragraph {para_num}: {paragraph.text[:100]}...")
            
            if len(current_page) >= 5:
                text_by_page.append('\n'.join(current_page))
                print(f"Created page with {len(current_page)} paragraphs")
                current_page = []
        
        if current_page:
            text_by_page.append('\n'.join(current_page))
            print(f"Added final page with {len(current_page)} paragraphs")
            
    except Exception as e:
        print(f"Error in DOC extraction: {str(e)}")
        import traceback
        print(traceback.format_exc())
    
    return text_by_page

def extract_document_data(file_path):
    data = []
    print(f"\nStarting extraction from: {file_path}")
    
    # Regular expressions for email and phone
    email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
    phone_pattern = r'''(?:\+\d{1,2}\s?)?\(?\d{3}\)?[\s.-]?\d{3}[\s.-]?\d{4}'''
    
    try:
        # Determine file type and extract text
        file_extension = os.path.splitext(file_path)[1].lower()
        print(f"File extension: {file_extension}")
        
        if file_extension == '.pdf':
            text_by_page = extract_text_from_pdf(file_path)
        elif file_extension in ['.doc', '.docx']:
            text_by_page = extract_text_from_doc(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        print(f"Number of pages extracted: {len(text_by_page)}")
        
        # Process each page individually
        for page_num, text in enumerate(text_by_page, 1):
            print(f"\nProcessing page {page_num}")
            
            # Extract candidate name from current page
            candidate_name = extract_candidate_name(text)
            if candidate_name:
                print(f"Found candidate name on page {page_num}: {candidate_name}")
            
            emails = re.findall(email_pattern, text)
            phones = re.findall(phone_pattern, text)
            
            print(f"Found emails: {emails}")
            print(f"Found phones: {phones}")
            
            if candidate_name or (emails and phones):
                person_data = {
                    'Name': candidate_name,
                    'Email': emails[0] if emails else '',
                    'Phone': phones[0] if phones else '',
                    'Found_On_Page': page_num  # page number reference
                }
                data.append(person_data)
                print(f"Added data: {person_data}")
            else:
                print("No data found on this page")
        
        # Post-process the data to combine information
        processed_data = []
        seen_names = set()
        current_data = {
            'Name': '', 
            'Email': '', 
            'Phone': '', 
            'Name_Found_On_Page': '',
            'Email_Found_On_Page': '',
            'Phone_Found_On_Page': ''
        }
        
        for entry in data:
            # If we find a name,we will use it
            if entry['Name'] and entry['Name'] not in seen_names:
                current_data['Name'] = entry['Name']
                current_data['Name_Found_On_Page'] = entry['Found_On_Page']
                seen_names.add(entry['Name'])
            
            # Adding email if we don't have one yet
            if entry['Email'] and not current_data['Email']:
                current_data['Email'] = entry['Email']
                current_data['Email_Found_On_Page'] = entry['Found_On_Page']
            
            # Adding phone if we don't have one yet
            if entry['Phone'] and not current_data['Phone']:
                current_data['Phone'] = entry['Phone']
                current_data['Phone_Found_On_Page'] = entry['Found_On_Page']
            
            # If we have all the data, adding it to processed_data
            if current_data['Name'] and (current_data['Email'] or current_data['Phone']):
                processed_data.append(current_data.copy())
                # Reset current_data but keep the name
                current_data = {
                    'Name': current_data['Name'],
                    'Email': '',
                    'Phone': '',
                    'Name_Found_On_Page': current_data['Name_Found_On_Page'],
                    'Email_Found_On_Page': '',
                    'Phone_Found_On_Page': ''
                }
        
        # Add any remaining data
        if current_data['Name'] and (current_data['Email'] or current_data['Phone']):
            processed_data.append(current_data)
        
        print(f"\nTotal data entries extracted: {len(processed_data)}")
        return processed_data
        
    except Exception as e:
        print(f"Error during extraction: {str(e)}")
        import traceback
        print(traceback.format_exc())
        return data

def extract_candidate_name(text):
    """
    Extract candidate name using multiple techniques
    """
    potential_names = set()
    
    # Name patterns with labels and formats
    name_patterns = [
        # Labels with different separators
        r'(?i)(?:name|candidate name|full name)\s*[:|-]\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)',
        r'(?i)(?:applicant name|candidate|profile)\s*[:|-]\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)',
        
        # Titles with names
        r'\b(?:Mr\.|Mrs\.|Ms\.|Dr\.|Prof\.|Er\.|Shri\.|Smt\.)\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)',
        
        # Standard name formats
        r'^([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,3})$',
        r'(^[A-Z]+ [A-Z]$)',
        r'(^[A-Z]+ [A-Z]+ [A-Z]$)',
        r'([A-Z][a-z]+\s+[A-Z]\.\s+[A-Z][a-z]+)'
        
        # Complex name formats
        r'([A-Z][a-z]+(?:\s+(?:van|de|der|den|das|dos|das|do|da|los|la|le|von|van)\s+[A-Z][a-z]+)+)',
        r'(?i)(?:s/o|d/o|w/o).*?([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)',
        r"([A-Z][a-z]+(?:['`]\s*[A-Z])?[a-z]*(?:\s+[A-Z][a-z]+)+)",
        r'([A-Z][a-z]+(?:-[A-Z][a-z]+)*(?:\s+[A-Z][a-z]+)+)'
    ]
    
    # Process first few lines
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    
    for line in lines[:7]:
        # Skip resume headers
        if any(header in line.lower() for header in ['resume', 'cv', 'curriculum vitae', 'biodata']):
            continue
        
        # Try regex patterns
        for pattern in name_patterns:
            matches = re.finditer(pattern, line)
            for match in matches:
                name = match.group(1).strip() if match.groups() else match.group(0).strip()
                if name and len(name.split()) >= 2:
                    potential_names.add(name)
        
        # Use NLP tools
        # SpaCy
        doc = nlp(line)
        for ent in doc.ents:
            if ent.label_ == "PERSON":
                potential_names.add(ent.text)
        
        # NLTK
        tokens = nltk.tokenize.word_tokenize(line)
        pos_tags = nltk.pos_tag(tokens)
        chunks = nltk.ne_chunk(pos_tags)
        for chunk in chunks:
            if hasattr(chunk, 'label') and chunk.label() == 'PERSON':
                name = ' '.join([c[0] for c in chunk.leaves()])
                potential_names.add(name)
    
    # Validate names
    validated_names = []
    for name in potential_names:
        parsed_name = HumanName(name)
        if parsed_name.first and parsed_name.last:
            cleaned_name = f"{parsed_name.first} {parsed_name.last}".strip()
            if is_valid_name(cleaned_name):
                validated_names.append(cleaned_name)
    
    # Sort by length and return the most complete name
    validated_names.sort(key=lambda x: len(x.split()), reverse=True)
    return validated_names[0] if validated_names else ""

def is_valid_name(name):
    """
    Validate if the extracted text is likely a real name
    """
    invalid_words = {
        'resume', 'cv', 'curriculum', 'vitae', 'email', 'phone', 'address',
        'education', 'experience', 'skills', 'objective', 'summary',
        'references', 'page', 'contact', 'profile', 'name', 'candidate',
        'applicant', 'date', 'birth', 'gender', 'nationality', 'status',
        'declaration', 'place', 'current', 'permanent', 'father', 'mother'
    }
    
    if len(name.split()) < 2 or len(name) < 4:
        return False
        
    if any(word.lower() in invalid_words for word in name.split()):
        return False
        
    if any(char.isdigit() for char in name):
        return False
        
    if not all(part[0].isupper() for part in name.split()):
        return False
    
    invalid_chars = set('!@#$%^&*()_+={}[]|\\:;"<>,.?/~`')
    if any(char in invalid_chars for char in name):
        return False
    
    valid_special_chars = set("-'")
    if any(char in name for char in valid_special_chars):
        for char in valid_special_chars:
            if char in name:
                parts = name.split(char)
                if not all(part.strip() and part.strip()[0].isupper() for part in parts):
                    return False
    
    return True

def save_to_excel(data, output_path):
    df = pd.DataFrame(data)
    
    # Reorder columns to show page numbers next to their corresponding fields
    column_order = [
        'Name', 'Name_Found_On_Page',
        'Email', 'Email_Found_On_Page',
        'Phone', 'Phone_Found_On_Page'
    ]
    
    # Ensuring all columns exist
    for col in column_order:
        if col not in df.columns:
            df[col] = ''
    
    # Reorder columns
    df = df[column_order]
    
    # Removeing duplicates 
    df = df.drop_duplicates(subset=['Email', 'Phone'], keep='first')
    
    # Removing rows where all fields are empty
    df = df.dropna(subset=['Name', 'Email', 'Phone'], how='all')
    
    # Sorting by Name
    df = df.sort_values('Name')
    
    # Save to Excel
    df.to_excel(output_path, index=False)
    print(f"Data saved to {output_path}")

def main():
    input_files = [
        "./pdf/merged_output.pdf",
        "./doc/input.docx"
    ]
    excel_path = "./output/output.xlsx"
    
    os.makedirs(os.path.dirname(excel_path), exist_ok=True)
    
    all_data = []
    
    for file_path in input_files:
        if os.path.exists(file_path):
            print(f"Processing {file_path}...")
            try:
                extracted_data = extract_document_data(file_path)
                all_data.extend(extracted_data)
            except Exception as e:
                print(f"Error processing {file_path}: {str(e)}")
        else:
            print(f"File not found: {file_path}")
    
    if all_data:
        save_to_excel(all_data, excel_path)
    else:
        print("No data was extracted from the input files.")

if __name__ == "__main__":
    main()